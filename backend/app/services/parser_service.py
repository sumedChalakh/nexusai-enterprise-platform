import io
import csv
import logging
from pathlib import Path

import fitz

log = logging.getLogger(__name__)


def parse_pdf(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            pages.append(f"[Page {i + 1}]\n{text}")
    doc.close()
    return "\n\n".join(pages)


def parse_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    # also pull table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def parse_txt(content: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


def parse_csv(content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


PARSERS = {
    "application/pdf": parse_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
    "text/plain": parse_txt,
    "text/csv": parse_csv,
}


def parse_document(content: bytes, content_type: str) -> str:
    parser = PARSERS.get(content_type)
    if not parser:
        raise ValueError(f"No parser for content type: {content_type}")
    try:
        text = parser(content)
        # Remove null bytes which PostgreSQL string/text columns reject
        text = text.replace("\x00", "")
        log.info("Parsed %d chars from %s", len(text), content_type)
        return text
    except Exception as e:
        log.error("Parsing failed (%s): %s", content_type, e)
        raise RuntimeError(f"Parse error: {e}") from e
